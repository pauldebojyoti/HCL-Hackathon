"""
Document Processing Module
Handles PDF parsing, text extraction, and chunking for RAG pipeline
"""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

# OCR imports for scanned PDFs
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    
    # Configure tesseract path for Windows
    import platform
    if platform.system() == "Windows":
        # Common Tesseract installation paths on Windows
        tesseract_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            r"C:\Users\Public\Tesseract-OCR\tesseract.exe"
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                break
    
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    pytesseract = None
    convert_from_path = None
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import docx

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Handles document processing including PDF parsing, text extraction, and chunking
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Initialize the document processor
        
        Args:
            chunk_size: Maximum size of text chunks in tokens
            chunk_overlap: Number of tokens to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF file with multiple fallback methods
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        if PdfReader is None:
            logger.error(f"No PDF library available. Install with: pip install pypdf2")
            return ""
            
        try:
            reader = PdfReader(file_path)
            text = ""
            
            for i, page in enumerate(reader.pages):
                try:
                    content = page.extract_text()
                    if content and content.strip():
                        text += content + "\n"
                    else:
                        logger.warning(f"No text extracted from page {i+1} of {file_path}")
                except Exception as page_error:
                    logger.warning(f"Error extracting page {i+1} from {file_path}: {str(page_error)}")
                    continue
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
                return text
            else:
                logger.warning(f"No text content found in PDF: {file_path}. Trying OCR...")
                # Try OCR as fallback for scanned PDFs
                return self.extract_pdf_text_with_ocr(file_path)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            # Try with different encoding
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        content = page.extract_text()
                        if content:
                            text += content + "\n"
                    if text.strip():
                        logger.info(f"Successfully extracted text with binary mode from {file_path}")
                        return text
            except Exception as e2:
                logger.error(f"Secondary PDF extraction failed for {file_path}: {str(e2)}")
            
            return ""
    
    def extract_pdf_text_with_ocr(self, file_path: str) -> str:
        """
        Extract text from PDF using OCR for scanned documents
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        if not OCR_AVAILABLE:
            logger.error("OCR libraries not available. Install with: pip install pytesseract pdf2image")
            return ""
            
        try:
            logger.info(f"Attempting OCR extraction from {file_path}")
            
            # Convert PDF to images
            images = convert_from_path(file_path)
            text = ""
            
            for i, image in enumerate(images):
                try:
                    # Extract text from image using OCR
                    page_text = pytesseract.image_to_string(image)
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                        logger.info(f"OCR extracted text from page {i+1}")
                    else:
                        logger.warning(f"No text extracted via OCR from page {i+1}")
                except Exception as page_error:
                    logger.warning(f"OCR failed for page {i+1}: {str(page_error)}")
                    continue
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters via OCR from {file_path}")
                return text
            else:
                logger.warning(f"No text content extracted via OCR from: {file_path}")
                return ""
                
        except Exception as e:
            logger.error(f"OCR extraction failed for {file_path}: {str(e)}")
            return ""
    
    def extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            logger.info(f"Successfully extracted text from {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    def extract_text_file(self, file_path: str) -> str:
        """
        Read text from TXT file
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content as string
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"Successfully read text from {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {str(e)}")
            return ""
    
    def process_document(self, file_path: str) -> str:
        """
        Process a document based on its file extension
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return self.extract_docx_text(file_path)
        elif file_extension == '.txt':
            return self.extract_text_file(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return ""
    
    def chunk_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Split text into chunks using LangChain text splitter
        
        Args:
            text: Text to be chunked
            metadata: Optional metadata to attach to chunks
            
        Returns:
            List of Document objects with chunked text
        """
        if not text.strip():
            return []
        
        # Create Document object
        doc = Document(page_content=text, metadata=metadata or {})
        
        # Split the document
        chunks = self.text_splitter.split_documents([doc])
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def process_documents_from_directory(self, directory_path: str) -> List[Document]:
        """
        Process all supported documents in a directory
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of Document chunks from all processed files
        """
        all_chunks = []
        supported_extensions = ['.pdf', '.txt', '.docx']
        
        directory = Path(directory_path)
        if not directory.exists():
            logger.error(f"Directory does not exist: {directory_path}")
            return all_chunks
        
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                logger.info(f"Processing file: {file_path}")
                
                # Extract text
                text = self.process_document(str(file_path))
                
                if text:
                    # Create metadata
                    metadata = {
                        'source': str(file_path),
                        'filename': file_path.name,
                        'file_type': file_path.suffix.lower()
                    }
                    
                    # Chunk the text
                    chunks = self.chunk_text(text, metadata)
                    all_chunks.extend(chunks)
        
        logger.info(f"Processed {len(all_chunks)} total chunks from directory")
        return all_chunks


def main():
    """
    Test the document processor
    """
    processor = DocumentProcessor()
    
    # Test processing a directory
    docs_directory = "data/docs"
    if os.path.exists(docs_directory):
        chunks = processor.process_documents_from_directory(docs_directory)
        print(f"Processed {len(chunks)} chunks from directory")
        
        if chunks:
            print(f"Sample chunk: {chunks[0].page_content[:200]}...")
    else:
        print(f"Directory {docs_directory} does not exist")


if __name__ == "__main__":
    main()